from __future__ import annotations

import io

from app.common.exception.code import StatusCode
from app.common.exception.exception import BusinessException
from app.infrastructure.log.logging_config import get_logger

logger = get_logger(__name__)


_SOURCE_TYPE_LABEL = {"report": "报告", "rag": "RAG", "web": "Web"}
_CITATION_STYLE_LABEL = {
    "direct": "直接引用",
    "paraphrase": "改写",
    "summary": "概括",
    "data": "数据引用",
}


def _extract_text_from_blocks(content_blocks: list[dict]) -> list[str]:
    """Extract displayable lines from content_blocks array.

    Inline [N] citation markers in the source text are preserved verbatim,
    so the rendered output keeps the linkage to the references list below.
    """
    lines: list[str] = []
    for block in content_blocks:
        btype = block.get("type", "")
        if btype in ("bullet_list", "numbered_list"):
            items = block.get("items") or []
            for item in items:
                if isinstance(item, str) and item.strip():
                    lines.append(item.strip())
        elif btype in ("paragraph", "quote", "subtitle"):
            text = block.get("text", "")
            if isinstance(text, str) and text.strip():
                lines.append(text.strip())
    return lines


def _format_reference_line(ref: dict, fallback_index: int) -> str:
    """Render one reference as `[N] source · 来源 · 引用方式`."""
    n = ref.get("ref_number")
    if not isinstance(n, int) or n < 1:
        n = fallback_index
    source = str(ref.get("source") or "").strip() or "(未知来源)"
    s_type_raw = str(ref.get("source_type") or "")
    c_style_raw = str(ref.get("citation_style") or "")
    s_type = _SOURCE_TYPE_LABEL.get(s_type_raw, s_type_raw)
    c_style = _CITATION_STYLE_LABEL.get(c_style_raw, c_style_raw)
    parts = [f"[{n}] {source}"]
    if s_type:
        parts.append(s_type)
    if c_style:
        parts.append(c_style)
    return " · ".join(parts)


class ExportService:
    """PPT 纯文本导出服务。"""

    def to_markdown(self, slide_content: dict) -> str:
        """
        将幻灯片 JSON 转换为 Markdown 字符串。

        支持新格式：
        {
          "presentation_title": "...",
          "slides": [
            {
              "slide_number": 1,
              "slide_type": "...",
              "title": "...",
              "content_blocks": [{"type": "bullet_list", "items": [...]}],
              "speaker_notes": "...",
              "references": [{"ref_number": 1, "source_type": "rag",
                              "source": "...", "citation_style": "direct",
                              "snippet": "..."}]
            }
          ]
        }
        """
        logger.info("Exporting slides to Markdown")
        slides: list[dict] = slide_content.get("slides", [])
        if not slides:
            logger.warning("Export to Markdown: slide list is empty")
            raise BusinessException.exc(StatusCode.EXPORT_FAILED.value)

        lines: list[str] = []
        presentation_title = slide_content.get("presentation_title", "")
        if presentation_title:
            lines.append(f"# {presentation_title}")
            lines.append("")

        for slide in slides:
            title = slide.get("title", "")
            content_blocks: list[dict] = slide.get("content_blocks") or []
            speaker_notes: str = slide.get("speaker_notes", "") or ""
            references: list[dict] = slide.get("references") or []

            if title:
                lines.append(f"## {title}")

            text_lines = _extract_text_from_blocks(content_blocks)
            for t in text_lines:
                lines.append(f"- {t}")

            if speaker_notes.strip():
                lines.append("")
                lines.append(f"> {speaker_notes.strip()}")

            if references:
                lines.append("")
                lines.append("**参考文献**")
                for i, ref in enumerate(references, start=1):
                    if isinstance(ref, dict):
                        lines.append(_format_reference_line(ref, i))

            lines.append("")

        result = "\n".join(lines).rstrip() + "\n"
        logger.info("Markdown export complete: %d slides, %d chars", len(slides), len(result))
        return result

    def to_word(self, slide_content: dict) -> bytes:
        """将幻灯片 JSON 转换为 Word (.docx) 字节流。"""
        logger.info("Exporting slides to Word (.docx)")
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
        except ImportError as e:
            logger.error("python-docx not installed: %s", e)
            raise BusinessException.exc(StatusCode.EXPORT_FAILED.value)

        slides: list[dict] = slide_content.get("slides", [])
        if not slides:
            logger.warning("Export to Word: slide list is empty")
            raise BusinessException.exc(StatusCode.EXPORT_FAILED.value)

        doc = Document()
        presentation_title = slide_content.get("presentation_title", "")
        if presentation_title:
            doc.add_heading(presentation_title, level=0)

        for slide in slides:
            title = slide.get("title", "")
            content_blocks: list[dict] = slide.get("content_blocks") or []
            speaker_notes: str = slide.get("speaker_notes", "") or ""
            references: list[dict] = slide.get("references") or []

            if title:
                doc.add_heading(title, level=2)

            text_lines = _extract_text_from_blocks(content_blocks)
            for t in text_lines:
                doc.add_paragraph(t, style="List Bullet")

            if speaker_notes.strip():
                para = doc.add_paragraph()
                run = para.add_run(f"备注：{speaker_notes.strip()}")
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

            if references:
                head = doc.add_paragraph()
                head_run = head.add_run("参考文献")
                head_run.bold = True
                head_run.font.size = Pt(10)
                for i, ref in enumerate(references, start=1):
                    if not isinstance(ref, dict):
                        continue
                    line = doc.add_paragraph()
                    run = line.add_run(_format_reference_line(ref, i))
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

        buf = io.BytesIO()
        doc.save(buf)
        data = buf.getvalue()
        logger.info("Word export complete: %d slides, %d bytes", len(slides), len(data))
        return data
